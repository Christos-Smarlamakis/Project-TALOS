# -*- coding: utf-8 -*-
#  Project TALOS
#  Copyright (C) 2026 Christos Smarlamakis
#
#  This program is free software: you can redistribute it and/or modify
#  it under the terms of the GNU Affero General Public License as
#  published by the Free Software Foundation, either version 3 of the
#  License, or (at your option) any later version.
#
#  For commercial licensing, please contact the author.

"""
Module: recommender.py (v4.1 - Structured Reports Update)
Project: TALOS v5.9.15

Description:
Αναβαθμισμένη έκδοση του Στρατηγικού Αναλυτή.
- Οι εξαγωγές (HTML, DOCX, MD) πλέον αντικατοπτρίζουν την ίδια δομή
  που εμφανίζεται στο terminal: Foundational → SOTA → Thematic Clusters.
- Προσθήκη operational_score σε όλες τις αναφορές.
- Φιλτράρισμα placeholder abstracts (DBLP "δεν παρέχει περίληψη") για καλύτερο clustering.
"""
import os
import sys
import os, sys
_P = os.path.abspath(os.path.dirname(__file__))
while _P and not os.path.exists(os.path.join(_P, 'talos.py')):
    _P = os.path.dirname(_P)
if _P: sys.path.insert(0, _P)
import sqlite3
from datetime import datetime
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Προσθέτουμε το root του project στο path
# Greek placeholder abstracts that skew clustering
NOISE_ABSTRACTS = [
    "Δεν υπάρχει διαθέσιμη περίληψη.",
    "Το DBLP δεν παρέχει περιλήψεις μέσω του API.",
    "Το Crossref δεν παρέχει πάντα περίληψη.",
    "Elsevier does not provide an abstract in this call.",
    "No abstract available.",
    "Abstract not provided by IEEE API.",
]

class ReadingRecommender:
    """
    Αναλύει τα άρθρα της βάσης δεδομένων του TALOS, εφαρμόζει μηχανική μάθηση
    (clustering) για να εντοπίσει θεματικές ενότητες, και προτείνει ένα
    στρατηγικό "μονοπάτι ανάγνωσης" για τον ερευνητή.
    """
    def __init__(self, db_name="talos_research.db"):
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        self.db_path = os.path.join(project_root, db_name)
        self.reports_dir = os.path.join(project_root, "data", "reports", "recommendations")
        os.makedirs(self.reports_dir, exist_ok=True)
        self.papers_df = self.load_papers_from_db()
        print(f"INFO: ReadingRecommender v4.1 αρχικοποιήθηκε με {len(self.papers_df)} άρθρα από τη βάση.")

    def load_papers_from_db(self) -> pd.DataFrame:
        try:
            with sqlite3.connect(self.db_path) as conn:
                query = """
                    SELECT id, doi, url, title, authors, publication_year, abstract, source,
                           strategic_score, operational_score, tactical_score, playground_score, overall_score,
                           evaluation_reasoning, evaluation_contribution, evaluation_utilization,
                           suggested_tags, processed_at
                    FROM papers
                """
                df = pd.read_sql_query(query, conn)
                df['processed_at'] = pd.to_datetime(df['processed_at'], errors='coerce')
                return df
        except Exception as e:
            print(f"FATAL: Αποτυχία φόρτωσης δεδομένων από τη βάση: {e}")
            return pd.DataFrame()

    def get_top_keywords_for_cluster(self, vectorizer, kmeans_model, cluster_id: int, top_n=4) -> str:
        cluster_center = kmeans_model.cluster_centers_[cluster_id]
        top_term_indices = cluster_center.argsort()[::-1][:top_n]
        feature_names = vectorizer.get_feature_names_out()
        keywords = [feature_names[i] for i in top_term_indices]
        return ", ".join(keywords)

    def _clean_abstract(self, text):
        """Αντικαθιστά placeholder abstracts με κενό για καλύτερο clustering."""
        if not isinstance(text, str):
            return ""
        for noise in NOISE_ABSTRACTS:
            if noise in text:
                return ""
        return text

    def run_analysis_and_reporting(self, num_clusters=5, min_score=7.0):
        print(f"\n[Βήμα 1/4] Φιλτράρισμα άρθρων με Overall Score >= {min_score}...")
        relevant_papers = self.papers_df[self.papers_df['overall_score'] >= min_score].copy()

        if len(relevant_papers) < num_clusters:
            print(f"WARNING: Δεν βρέθηκαν αρκετά σχετικά άρθρα ({len(relevant_papers)}) για να δημιουργηθούν {num_clusters} clusters. Τερματισμός.")
            return

        print(f"Βρέθηκαν {len(relevant_papers)} σχετικά άρθρα για περαιτέρω ανάλυση.")

        print("\n[Βήμα 2/4] Ομαδοποίηση άρθρων σε θεματικές ενότητες (Clustering)...")
        # Clean abstracts for better clustering
        relevant_papers['clean_abstract'] = relevant_papers['abstract'].apply(self._clean_abstract)
        corpus = (relevant_papers['title'] + ' ' + relevant_papers['clean_abstract']).fillna('')
        vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
        X = vectorizer.fit_transform(corpus)
        kmeans = KMeans(n_clusters=num_clusters, random_state=42, n_init='auto')
        relevant_papers['cluster'] = kmeans.fit_predict(X)
        print("Η ομαδοποίηση ολοκληρώθηκε.")

        print("\n[Βήμα 3/4] Δημιουργία στρατηγικού μονοπατιού ανάγνωσης...")
        # Foundational: oldest papers with high scores first (≥ 7.0), then fall back to ≥ 5.0
        found_elite = relevant_papers[relevant_papers['overall_score'] >= 7.0].sort_values(by='publication_year', ascending=True)
        if len(found_elite) < 3:
            found_elite = relevant_papers[relevant_papers['overall_score'] >= 5.0].sort_values(by='publication_year', ascending=True)
        foundational_papers = found_elite.head(5)
        # State-of-the-Art: newest high-scoring papers by publication_year
        hot_elite = relevant_papers[relevant_papers['overall_score'] >= 7.0].sort_values(by='publication_year', ascending=False)
        if len(hot_elite) < 3:
            hot_elite = relevant_papers[relevant_papers['overall_score'] >= 5.0].sort_values(by='publication_year', ascending=False)
        hot_papers = hot_elite.head(5)

        # Build cluster data
        clusters_data = []
        for i in range(num_clusters):
            cluster_df = relevant_papers[relevant_papers['cluster'] == i]
            if len(cluster_df) < 2:
                continue
            keywords = self.get_top_keywords_for_cluster(vectorizer, kmeans, i)
            top_papers = cluster_df.sort_values(by='overall_score', ascending=False).head(3)
            clusters_data.append({
                'keywords': keywords,
                'papers': top_papers
            })

        # Print to terminal
        self._print_structured_report(foundational_papers, hot_papers, clusters_data)

        print("\n[Βήμα 4/4] Εξαγωγή πλήρων αναφορών...")
        self.export_structured_reports(foundational_papers, hot_papers, clusters_data, relevant_papers)
        print("\nΗ διαδικασία ολοκληρώθηκε.")

    def _print_structured_report(self, foundational, hot, clusters):
        print("\n" + "="*80)
        print("      ** ΣΤΡΑΤΗΓΙΚΟ ΜΟΝΟΠΑΤΙ ΑΝΑΓΝΩΣΗΣ (Σύνοψη) **")
        print("="*80)
        print("\n--- 📚 1. ΘΕΜΕΛΙΩΔΗ ΑΡΘΡΑ (Foundational Papers) ---\n")
        for _, row in foundational.iterrows():
            scores = f"S:{int(row['strategic_score'])} O:{int(row['operational_score'])} T:{int(row['tactical_score'])} P:{int(row['playground_score'])}"
            print(f"  [{row['overall_score']:.1f}] {row['title']} ({int(row['publication_year'])}) — {scores}")

        print("\n--- 🔥 2. ΤΕΛΕΥΤΑΙΕΣ ΕΞΕΛΙΞΕΙΣ (State-of-the-Art) ---\n")
        for _, row in hot.iterrows():
            scores = f"S:{int(row['strategic_score'])} O:{int(row['operational_score'])} T:{int(row['tactical_score'])} P:{int(row['playground_score'])}"
            print(f"  [{row['overall_score']:.1f}] {row['title']} ({int(row['publication_year'])}) — {scores}")

        print("\n--- 🔬 3. ΕΞΕΙΔΙΚΕΥΜΕΝΑ ΘΕΜΑΤΑ (Thematic Clusters) ---\n")
        for i, cluster in enumerate(clusters):
            print(f"---  Cluster {i+1}: Εστίαση σε [ {cluster['keywords']} ] ---")
            for _, row in cluster['papers'].iterrows():
                scores = f"S:{int(row['strategic_score'])} O:{int(row['operational_score'])} T:{int(row['tactical_score'])} P:{int(row['playground_score'])}"
                print(f"  [{row['overall_score']:.1f}] {row['title']} ({int(row['publication_year'])}) — {scores}")
        print("="*80)

    def _paper_to_dict(self, row):
        """Convert a DataFrame row to a safe dict for templating."""
        return {
            'title': str(row.get('title', 'N/A')),
            'authors': str(row.get('authors', 'N/A')),
            'source': str(row.get('source', 'N/A')),
            'year': int(row['publication_year']) if pd.notna(row.get('publication_year')) else 'N/A',
            'overall': float(row.get('overall_score', 0)),
            'strategic': int(row.get('strategic_score', 0)),
            'operational': int(row.get('operational_score', 0)),
            'tactical': int(row.get('tactical_score', 0)),
            'playground': int(row.get('playground_score', 0)),
            'doi': str(row.get('doi', '')),
            'url': str(row.get('url', '#')),
            'abstract': str(row.get('abstract', ''))[:500],
            'reasoning': str(row.get('evaluation_reasoning', 'N/A')),
            'contribution': str(row.get('evaluation_contribution', 'N/A')),
            'utilization': str(row.get('evaluation_utilization', 'N/A')),
            'tags': str(row.get('suggested_tags', 'N/A')),
        }

    def export_structured_reports(self, foundational, hot, clusters, all_papers):
        """Exports all 3 formats with the same structured layout."""
        # Convert DataFrames to lists of dicts
        found_list = [self._paper_to_dict(row) for _, row in foundational.iterrows()]
        hot_list = [self._paper_to_dict(row) for _, row in hot.iterrows()]
        clusters_list = []
        for c in clusters:
            clusters_list.append({
                'keywords': c['keywords'],
                'papers': [self._paper_to_dict(row) for _, row in c['papers'].iterrows()]
            })
        # Top 50 overall for the full sorted list
        top50 = [self._paper_to_dict(row) for _, row in all_papers.sort_values(by='overall_score', ascending=False).head(50).iterrows()]

        timestamp = datetime.now().strftime("%Y%m%d")
        self._export_structured_html(found_list, hot_list, clusters_list, top50, timestamp)
        self._export_structured_docx(found_list, hot_list, clusters_list, top50, timestamp)
        self._export_structured_markdown(found_list, hot_list, clusters_list, top50, timestamp)

    def _export_structured_html(self, found, hot, clusters, top50, timestamp):
        """Interactive HTML with tabbed sections matching terminal structure."""
        filename = os.path.join(self.reports_dir, f"talos_strategic_report_{timestamp}.html")

        def paper_card(p):
            scores = f"S:{p['strategic']} / O:{p['operational']} / T:{p['tactical']} / P:{p['playground']}"
            link = p['doi'] if p['doi'] else p['url']
            link_html = f'<a href="{p["url"]}" target="_blank">{link}</a>' if p['url'] and p['url'] != '#' else 'N/A'
            return f"""
            <div class="paper-card">
                <div class="paper-score">[{p['overall']:.1f}]</div>
                <div class="paper-content">
                    <h3>{p['title']}</h3>
                    <div class="paper-meta">
                        <span>👤 {p['authors']}</span> |
                        <span>📚 {p['source']} ({p['year']})</span> |
                        <span>🔗 {link_html}</span>
                    </div>
                    <div class="paper-scores">{scores}</div>
                    <div class="paper-reasoning">{p['reasoning']}</div>
                </div>
            </div>"""

        found_html = "\n".join(paper_card(p) for p in found)
        hot_html = "\n".join(paper_card(p) for p in hot)

        clusters_html = ""
        for i, c in enumerate(clusters):
            cards = "\n".join(paper_card(p) for p in c['papers'])
            clusters_html += f"""
            <div class="cluster-section">
                <h3>🔬 Cluster {i+1}: Εστίαση σε [ {c['keywords']} ]</h3>
                <div class="papers-grid">{cards}</div>
            </div>"""

        # Top 50 table
        top50_rows = ""
        for p in top50:
            link = p['doi'] if p['doi'] else p['url']
            top50_rows += f"""
            <tr>
                <td>{p['overall']:.1f}</td>
                <td>{p['strategic']}</td>
                <td>{p['operational']}</td>
                <td>{p['tactical']}</td>
                <td>{p['playground']}</td>
                <td>{p['title']}</td>
                <td>{p['authors']}</td>
                <td>{p['source']}</td>
                <td>{p['year']}</td>
                <td><a href="{p['url']}" target="_blank">Open</a></td>
            </tr>"""

        html = f"""<!DOCTYPE html>
<html lang="el">
<head>
    <meta charset="UTF-8">
    <title>TALOS Strategic Reading Report</title>
    <link rel="stylesheet" type="text/css" href="https://cdn.datatables.net/1.13.6/css/jquery.dataTables.css">
    <script src="https://code.jquery.com/jquery-3.7.0.js"></script>
    <script src="https://cdn.datatables.net/1.13.6/js/jquery.dataTables.js"></script>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 20px; background: #f4f6f8; color: #333; max-width: 1200px; margin: auto; }}
        h1 {{ color: #2c3e50; text-align: center; }}
        h2 {{ color: #2c3e50; border-bottom: 3px solid #27ae60; padding-bottom: 10px; margin-top: 40px; }}
        h3 {{ color: #34495e; }}
        .tabs {{ display: flex; gap: 5px; margin: 20px 0; }}
        .tab-btn {{ padding: 10px 20px; background: #eee; border: none; cursor: pointer; border-radius: 5px 5px 0 0; font-size: 14px; }}
        .tab-btn.active {{ background: #27ae60; color: white; }}
        .tab-content {{ display: none; }}
        .tab-content.active {{ display: block; }}
        .paper-card {{ display: flex; gap: 15px; background: white; padding: 15px; margin: 10px 0; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
        .paper-score {{ font-size: 28px; font-weight: bold; color: #27ae60; min-width: 60px; text-align: center; }}
        .paper-content {{ flex: 1; }}
        .paper-content h3 {{ margin: 0 0 5px 0; font-size: 16px; }}
        .paper-meta {{ font-size: 12px; color: #777; margin-bottom: 5px; }}
        .paper-scores {{ font-size: 12px; color: #555; margin-bottom: 5px; font-family: monospace; }}
        .paper-reasoning {{ font-size: 13px; color: #555; line-height: 1.5; }}
        .cluster-section {{ margin: 20px 0; }}
        #reportTable_wrapper {{ margin-top: 20px; }}
    </style>
</head>
<body>
    <h1>📋 TALOS — Στρατηγική Αναφορά Ανάγνωσης</h1>
    <p style="text-align:center; color: #777;">{datetime.now().strftime('%d-%m-%Y %H:%M')} | {len(top50)} papers analyzed</p>

    <div class="tabs">
        <button class="tab-btn active" onclick="showTab('structured')">📖 Structured Path</button>
        <button class="tab-btn" onclick="showTab('top50')">📊 Top 50 (Interactive Table)</button>
    </div>

    <div class="tab-content active" id="tab-structured">
        <h2>📚 1. Θεμελιώδη Άρθρα (Foundational Papers)</h2>
        <p style="color:#777;">Ξεκινήστε εδώ για να κατανοήσετε τις βάσεις του πεδίου.</p>
        {found_html}

        <h2>🔬 2. Εξειδικευμένα Θέματα (Thematic Clusters)</h2>
        <p style="color:#777;">Εμβαθύνετε σε συγκεκριμένες θεματικές περιοχές.</p>
        {clusters_html}

        <h2>🔥 3. Τελευταίες Εξελίξεις (State-of-the-Art)</h2>
        <p style="color:#777;">Ολοκληρώστε με τα πιο πρόσφατα ευρήματα.</p>
        {hot_html}
    </div>

    <div class="tab-content" id="tab-top50">
        <h2>📊 Top 50 Papers by Overall Score</h2>
        <table id="reportTable" class="display">
            <thead>
                <tr>
                    <th>Overall</th><th>S</th><th>O</th><th>T</th><th>P</th>
                    <th>Title</th><th>Authors</th><th>Source</th><th>Year</th><th>Link</th>
                </tr>
            </thead>
            <tbody>{top50_rows}</tbody>
        </table>
    </div>

    <script>
        function showTab(tabName) {{
            document.querySelectorAll('.tab-content').forEach(t => t.classList.remove('active'));
            document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
            document.getElementById('tab-' + tabName).classList.add('active');
            event.target.classList.add('active');
            if (tabName === 'top50') {{
                if (!window.tableInitialized) {{
                    $('#reportTable').DataTable({{ pageLength: 25, order: [[0, "desc"]] }});
                    window.tableInitialized = true;
                }}
            }}
        }}
    </script>
</body>
</html>"""

        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"SUCCESS: Η δομημένη αναφορά HTML αποθηκεύτηκε στο: {filename}")

    def _export_structured_docx(self, found, hot, clusters, top50, timestamp):
        """Structured DOCX matching terminal layout."""
        filename = os.path.join(self.reports_dir, f"talos_strategic_reading_path_{timestamp}.docx")
        doc = Document()

        # Title
        title = doc.add_heading('TALOS — Στρατηγικό Μονοπάτι Ανάγνωσης', 0)
        doc.add_paragraph(f"Αναφορά: {datetime.now().strftime('%d-%m-%Y %H:%M')} | {len(top50)} papers")

        def add_paper_section(doc, p, show_reasoning=True):
            doc.add_heading(f"[{p['overall']:.1f}] {p['title']}", level=2)
            para = doc.add_paragraph()
            para.add_run('Authors: ').bold = True
            para.add_run(p['authors'])
            para = doc.add_paragraph()
            para.add_run('Source: ').bold = True
            para.add_run(f"{p['source']} ({p['year']})")
            para = doc.add_paragraph()
            para.add_run('Scores: ').bold = True
            para.add_run(f"S:{p['strategic']} O:{p['operational']} T:{p['tactical']} P:{p['playground']}")
            para = doc.add_paragraph()
            para.add_run('Link: ').bold = True
            para.add_run(p['url'] if p['url'] != '#' else 'N/A')
            if show_reasoning and p['reasoning'] and p['reasoning'] != 'N/A':
                doc.add_heading('AI Analysis', level=3)
                doc.add_paragraph(p['reasoning'], style='Intense Quote')
            doc.add_paragraph()

        # Section 1: Foundational
        doc.add_heading('📚 1. Θεμελιώδη Άρθρα (Foundational Papers)', level=1)
        doc.add_paragraph('Ξεκινήστε εδώ για να κατανοήσετε τις βάσεις του πεδίου.')
        for p in found:
            add_paper_section(doc, p)

        # Section 2: Thematic Clusters
        doc.add_heading('🔬 2. Εξειδικευμένα Θέματα (Thematic Clusters)', level=1)
        doc.add_paragraph('Εμβαθύνετε σε συγκεκριμένες θεματικές περιοχές.')
        for i, c in enumerate(clusters):
            doc.add_heading(f"Cluster {i+1}: {c['keywords']}", level=2)
            for p in c['papers']:
                add_paper_section(doc, p)

        # Section 3: State-of-the-Art
        doc.add_heading('🔥 3. Τελευταίες Εξελίξεις (State-of-the-Art)', level=1)
        doc.add_paragraph('Ολοκληρώστε με τα πιο πρόσφατα ευρήματα.')
        for p in hot:
            add_paper_section(doc, p)

        doc.save(filename)
        print(f"SUCCESS: Η δομημένη αναφορά DOCX αποθηκεύτηκε στο: {filename}")

    def _export_structured_markdown(self, found, hot, clusters, top50, timestamp):
        """Structured Markdown matching terminal layout."""
        filename = os.path.join(self.reports_dir, f"talos_strategic_reading_path_{timestamp}.md")

        lines = [
            f"# 📋 TALOS — Στρατηγικό Μονοπάτι Ανάγνωσης",
            f"_Αναφορά: {datetime.now().strftime('%d-%m-%Y %H:%M')} | {len(top50)} papers_\n",
        ]

        def format_paper(p):
            scores = f"S:{p['strategic']} O:{p['operational']} T:{p['tactical']} P:{p['playground']}"
            link = p['doi'] if p['doi'] else p['url']
            result = [
                f"### [{p['overall']:.1f}] {p['title']}",
                f"- **Authors:** {p['authors']}",
                f"- **Source:** {p['source']} ({p['year']})",
                f"- **Scores:** {scores}",
                f"- **Link:** [{link}]({p['url']})",
            ]
            if p['reasoning'] and p['reasoning'] != 'N/A':
                result.append(f"\n> 💡 {p['reasoning']}\n")
            return "\n".join(result)

        # Section 1
        lines.append("## 📚 1. Θεμελιώδη Άρθρα (Foundational Papers)")
        lines.append("_Ξεκινήστε εδώ για να κατανοήσετε τις βάσεις του πεδίου._\n")
        for p in found:
            lines.append(format_paper(p))
            lines.append("")

        # Section 2
        lines.append("## 🔬 2. Εξειδικευμένα Θέματα (Thematic Clusters)")
        lines.append("_Εμβαθύνετε σε συγκεκριμένες θεματικές περιοχές._\n")
        for i, c in enumerate(clusters):
            lines.append(f"### Cluster {i+1}: {c['keywords']}\n")
            for p in c['papers']:
                lines.append(format_paper(p))
                lines.append("")

        # Section 3
        lines.append("## 🔥 3. Τελευταίες Εξελίξεις (State-of-the-Art)")
        lines.append("_Ολοκληρώστε με τα πιο πρόσφατα ευρήματα._\n")
        for p in hot:
            lines.append(format_paper(p))
            lines.append("")

        # Top 50 quick list
        lines.append("---\n")
        lines.append("## 📊 Top 50 by Overall Score\n")
        lines.append("| # | Score | S | O | T | P | Title | Source | Year |")
        lines.append("|---|-------|---|---|---|---|-------|--------|------|")
        for i, p in enumerate(top50):
            title_short = p['title'][:60] + '...' if len(p['title']) > 60 else p['title']
            lines.append(f"| {i+1} | {p['overall']:.1f} | {p['strategic']} | {p['operational']} | {p['tactical']} | {p['playground']} | {title_short} | {p['source']} | {p['year']} |")

        with open(filename, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        print(f"SUCCESS: Η δομημένη αναφορά Markdown αποθηκεύτηκε στο: {filename}")


if __name__ == "__main__":
    recommender = ReadingRecommender()
    recommender.run_analysis_and_reporting()