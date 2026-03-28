# -*- coding: utf-8 -*-
#  Project TALOS
#  Copyright (C) 2026 Christos Smarlamakis
#
#  This program is free software: you can redistribute it and/or modify
#  it under the terms of the GNU Affero General Public License as
#  published by the Free Software Foundation, either version 3 of the
#  License, or (at your option) any later version.
#
#  For commercial licensing, please contact the author.

"""
Module: recommender.py (v4.0 - JSON Reliability Update)
Project: TALOS v2.21.0

Description:
Η πλήρως αναβαθμισμένη έκδοση του Στρατηγικού Αναλυτή, εναρμονισμένη με την
αρχιτεκτονική της "Αναβάθμισης Αξιοπιστίας".
- Διαβάζει το νέο, εμπλουτισμένο σχήμα της βάσης δεδομένων, συμπεριλαμβανομένων
  όλων των δομημένων πεδίων αξιολόγησης (reasoning, contribution, etc.).
- Η λογική του clustering παραμένει η ίδια, αλλά οι τελικές αναφορές που παράγονται
  (HTML, DOCX, MD) είναι πλέον πολύ πιο πλούσιες και καθαρές, καθώς αξιοποιούν
  τα νέα, διακριτά πεδία δεδομένων.
"""
import os
import sys
import sqlite3
from datetime import datetime
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from docx import Document

# Προσθέτουμε το root του project στο path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

class ReadingRecommender:
    """
    Αναλύει τα άρθρα της βάσης δεδομένων του TALOS, εφαρμόζει μηχανική μάθηση
    (clustering) για να εντοπίσει θεματικές ενότητες, και προτείνει ένα
    στρατηγικό "μονοπάτι ανάγνωσης" για τον ερευνητή.
    """
    def __init__(self, db_name="talos_research.db"):
        """
        Αρχικοποιεί τον Recommender.

        Args:
            db_name (str): Το όνομα του αρχείου της βάσης δεδομένων SQLite.
        """
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        self.db_path = os.path.join(project_root, db_name)
        # Ορίζουμε τον φάκελο όπου θα αποθηκευτούν οι αναφορές
        self.reports_dir = os.path.join(project_root, "reports", "recommendations")
        os.makedirs(self.reports_dir, exist_ok=True)
        # Φορτώνουμε τα δεδομένα από τη βάση κατά την αρχικοποίηση
        self.papers_df = self.load_papers_from_db()
        print(f"INFO: ReadingRecommender v4.0 αρχικοποιήθηκε με {len(self.papers_df)} άρθρα από τη βάση.")

    def load_papers_from_db(self) -> pd.DataFrame:
        """
        Φορτώνει όλα τα σχετικά δεδομένα από τη βάση δεδομένων σε ένα Pandas DataFrame.

        Returns:
            pd.DataFrame: Ένα DataFrame που περιέχει τα δεδομένα των άρθρων.
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                # **ΑΝΑΒΑΘΜΙΣΗ**: Το query επιλέγει πλέον όλα τα νέα, δομημένα πεδία.
                query = """
                    SELECT id, doi, url, title, authors, publication_year, abstract, source,
                           strategic_score, tactical_score, playground_score, overall_score,
                           evaluation_reasoning, evaluation_contribution, evaluation_utilization,
                           suggested_tags, processed_at
                    FROM papers
                """
                df = pd.read_sql_query(query, conn)
                # Μετατρέπουμε τη στήλη ημερομηνίας σε αντικείμενο datetime για σωστή ταξινόμηση
                df['processed_at'] = pd.to_datetime(df['processed_at'], errors='coerce')
                return df
        except Exception as e:
            print(f"FATAL: Αποτυχία φόρτωσης δεδομένων από τη βάση: {e}")
            return pd.DataFrame() # Επιστρέφουμε ένα κενό DataFrame σε περίπτωση σφάλματος

    def get_top_keywords_for_cluster(self, vectorizer, kmeans_model, cluster_id: int, top_n=4) -> str:
        """
        Βρίσκει τις πιο σημαντικές λέξεις-κλειδιά για ένα συγκεκριμένο cluster.

        Args:
            vectorizer: Το εκπαιδευμένο TfidfVectorizer.
            kmeans_model: Το εκπαιδευμένο KMeans model.
            cluster_id (int): Το ID του cluster προς ανάλυση.
            top_n (int): Ο αριθμός των λέξεων-κλειδιών που θα επιστραφούν.

        Returns:
            str: Ένα string με τις λέξεις-κλειδιά, χωρισμένες με κόμμα.
        """
        # Παίρνουμε το κέντρο του cluster και τις top λέξεις-κλειδιά
        cluster_center = kmeans_model.cluster_centers_[cluster_id]
        top_term_indices = cluster_center.argsort()[::-1][:top_n]
        feature_names = vectorizer.get_feature_names_out()
        keywords = [feature_names[i] for i in top_term_indices]
        return ", ".join(keywords)

    def run_analysis_and_reporting(self, num_clusters=5, min_score=4.0):
        """
        Ενορχηστρώνει την πλήρη ροή της ανάλυσης: φιλτράρισμα, clustering,
        δημιουργία μονοπατιού ανάγνωσης και εξαγωγή αναφορών.
        """
        print(f"\n[Βήμα 1/4] Φιλτράρισμα άρθρων με Overall Score >= {min_score}...")
        relevant_papers = self.papers_df[self.papers_df['overall_score'] >= min_score].copy()
        
        if len(relevant_papers) < num_clusters:
            print(f"WARNING: Δεν βρέθηκαν αρκετά σχετικά άρθρα ({len(relevant_papers)}) για να δημιουργηθούν {num_clusters} clusters. Τερματισμός.")
            return
        
        print(f"Βρέθηκαν {len(relevant_papers)} σχετικά άρθρα για περαιτέρω ανάλυση.")

        print("\n[Βήμα 2/4] Ομαδοποίηση άρθρων σε θεματικές ενότητες (Clustering)...")
        # Δημιουργούμε το "σώμα" κειμένου για κάθε άρθρο, συνδυάζοντας τίτλο και περίληψη
        corpus = (relevant_papers['title'] + ' ' + relevant_papers['abstract']).fillna('')
        vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
        X = vectorizer.fit_transform(corpus)
        kmeans = KMeans(n_clusters=num_clusters, random_state=42, n_init='auto')
        relevant_papers['cluster'] = kmeans.fit_predict(X)
        print("Η ομαδοποίηση ολοκληρώθηκε.")

        print("\n[Βήμα 3/4] Δημιουργία στρατηγικού μονοπατιού ανάγνωσης...")
        # Βρίσκουμε τα παλαιότερα άρθρα υψηλής σημασίας (θεμελιώδη)
        foundational_papers = relevant_papers.sort_values(by='publication_year', ascending=True).head(5)
        # Βρίσκουμε τα νεότερα άρθρα υψηλής σημασίας (state-of-the-art)
        hot_papers = relevant_papers.sort_values(by='processed_at', ascending=False).head(5)

        # Εκτύπωση της σύνοψης στο τερματικό
        print("\n" + "="*80)
        print("      ** ΣΤΡΑΤΗΓΙΚΟ ΜΟΝΟΠΑΤΙ ΑΝΑΓΝΩΣΗΣ (Σύνοψη) **")
        print("="*80)
        print("\n--- 📚 1. ΘΕΜΕΛΙΩΔΗ ΑΡΘΡΑ (Foundational Papers) ---\n")
        for _, row in foundational_papers.iterrows(): print(f"  [{row['overall_score']:.1f}] {row['title']} ({row['publication_year']})")
        
        print("\n--- 🔥 2. ΤΕΛΕΥΤΑΙΕΣ ΕΞΕΛΙΞΕΙΣ (State-of-the-Art) ---\n")
        for _, row in hot_papers.iterrows(): print(f"  [{row['overall_score']:.1f}] {row['title']} ({row['publication_year']})")
        
        print("\n--- 🔬 3. ΕΞΕΙΔΙΚΕΥΜΕΝΑ ΘΕΜΑΤΑ (Thematic Clusters) ---\n")
        for i in range(num_clusters):
            cluster_df = relevant_papers[relevant_papers['cluster'] == i]
            if len(cluster_df) < 2: continue
            keywords = self.get_top_keywords_for_cluster(vectorizer, kmeans, i)
            print(f"---  Cluster {i+1}: Εστίαση σε [ {keywords} ] ---")
            for _, row in cluster_df.sort_values(by='overall_score', ascending=False).head(3).iterrows():
                print(f"  [{row['overall_score']:.1f}] {row['title']} ({row['publication_year']})")
        print("="*80)

        print("\n[Βήμα 4/4] Εξαγωγή πλήρων αναφορών...")
        self.generate_and_export_reports(self.papers_df.sort_values(by='overall_score', ascending=False))
        print("\nΗ διαδικασία ολοκληρώθηκε.")

    def generate_and_export_reports(self, df: pd.DataFrame):
        """
        Προετοιμάζει το DataFrame και καλεί τις επιμέρους μεθόδους εξαγωγής
        για κάθε τύπο αναφοράς (CSV, HTML, DOCX, Markdown).
        """
        # **ΑΝΑΒΑΘΜΙΣΗ**: Χρησιμοποιούμε τα νέα, καθαρά πεδία για τις αναφορές
        df_export = df.rename(columns={
            'overall_score': 'Overall', 'tactical_score': 'Tactical',
            'strategic_score': 'Strategic', 'playground_score': 'Simulation',
            'title': 'Title', 'authors': 'Authors', 'source': 'Source',
            'processed_at': 'Date', 'url': 'Link', 'abstract': 'Abstract',
            'evaluation_reasoning': 'AI Reasoning', 'suggested_tags': 'AI Tags'
        })
        
        # Επιλέγουμε τις στήλες που θέλουμε και τη σειρά τους
        final_columns = [
            'Overall', 'Tactical', 'Strategic', 'Simulation', 'Title', 'Authors',
            'publication_year', 'Source', 'Date', 'Link', 'AI Reasoning', 'AI Tags'
        ]
        df_export = df_export[final_columns]
        
        # Καλούμε κάθε μέθοδο εξαγωγής
        self.export_to_html(df_export)
        self.export_to_docx(df_export.head(50)) # DOCX μόνο για τα top 50
        self.export_to_markdown(df_export.head(50)) # MD μόνο για τα top 50

    def export_to_html(self, df: pd.DataFrame):
        """Εξάγει μια πλήρως διαδραστική αναφορά HTML με δυνατότητες ταξινόμησης."""
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = os.path.join(self.reports_dir, f"talos_strategic_report_{timestamp}.html")
        
        df_html = df.copy()
        df_html['Link'] = df_html['Link'].apply(lambda x: f'<a href="{x}" target="_blank">Open</a>' if x else 'N/A')
        df_html['Overall'] = df_html['Overall'].apply(lambda x: f"{x:.2f}")
        
        html_table = df_html.to_html(escape=False, index=False, table_id="reportTable", classes="display")
        
        html_template = f"""
        <!DOCTYPE html><html lang="el"><head>
            <meta charset="UTF-8"><title>TALOS Strategic Report</title>
            <link rel="stylesheet" type="text/css" href="https://cdn.datatables.net/1.13.6/css/jquery.dataTables.css">
            <script src="https://code.jquery.com/jquery-3.7.0.js"></script>
            <script src="https://cdn.datatables.net/1.13.6/js/jquery.dataTables.js"></script>
        </head><body>
            <h1>TALOS - Στρατηγική Αναφορά Ανάγνωσης</h1>{html_table}
            <script>
                $(document).ready(function() {{ $('#reportTable').DataTable({{ "pageLength": 25, "order": [[0, "desc"]] }}); }});
            </script>
        </body></html>
        """
        with open(filename, 'w', encoding='utf-8') as f: f.write(html_template)
        print(f"SUCCESS: Η διαδραστική αναφορά HTML αποθηκεύτηκε στο: {filename}")

    def export_to_docx(self, df: pd.DataFrame):
        """Εξάγει μια συνοπτική αναφορά σε μορφή DOCX για τα πιο σημαντικά άρθρα."""
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = os.path.join(self.reports_dir, f"talos_summary_{timestamp}.docx")
        doc = Document()
        doc.add_heading('TALOS Summary Report', 0)
        doc.add_paragraph(f"Αναφορά που δημιουργήθηκε στις: {datetime.now().strftime('%d-%m-%Y')}")
        
        for _, row in df.iterrows():
            doc.add_heading(f"[{row['Overall']:.2f}] {row['Title']}", level=2)
            p = doc.add_paragraph(); p.add_run('Authors: ').bold = True; p.add_run(str(row['Authors']))
            p = doc.add_paragraph(); p.add_run('Source: ').bold = True; p.add_run(f"{row['Source']} ({row['publication_year']})")
            p = doc.add_paragraph(); p.add_run('Link: ').bold = True; p.add_run(row['Link'])
            doc.add_heading("AI Analysis", level=3)
            doc.add_paragraph(str(row['AI Reasoning']), style='Intense Quote')
            doc.add_paragraph() # Κενή γραμμή
            
        doc.save(filename)
        print(f"SUCCESS: Η συνοπτική αναφορά DOCX αποθηκεύτηκε στο: {filename}")

    def export_to_markdown(self, df: pd.DataFrame):
        """Εξάγει μια συνοπτική αναφορά σε μορφή Markdown."""
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = os.path.join(self.reports_dir, f"talos_summary_{timestamp}.md")
        
        content = [f"# TALOS Summary Report - {datetime.now().strftime('%d-%m-%Y')}"]
        for _, row in df.iterrows():
            content.append(f"\n---\n")
            content.append(f"## [{row['Overall']:.2f}] {row['Title']}")
            content.append(f"- **Authors:** {row['Authors']}")
            content.append(f"- **Source:** {row['Source']} ({row['publication_year']})")
            content.append(f"- **Link:** [{row['Link']}]({row['Link']})")
            content.append(f"\n> [!info] TALOS Analysis\n> {row['AI Reasoning']}\n")

        with open(filename, 'w', encoding='utf-8') as f: f.write("\n".join(content))
        print(f"SUCCESS: Η συνοπτική αναφορά Markdown αποθηκεύτηκε στο: {filename}")

if __name__ == "__main__":
    recommender = ReadingRecommender()
    recommender.run_analysis_and_reporting()